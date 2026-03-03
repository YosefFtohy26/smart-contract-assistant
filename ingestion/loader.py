import fitz
from docx import Document as DocxDocument
from pathlib import Path
from langchain_core.documents import Document
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoadError(Exception):
    pass


def extract_text_from_pdf(file_path: Path):
    documents = []
    try:
        with fitz.open(file_path) as doc:
            for page_number, page in enumerate(doc, start=1):
                text = page.get_text()
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file_path.name,
                            "page": page_number
                        }
                    )
                )
        return documents
    except Exception as e:
        raise DocumentLoadError(f"PDF extraction failed: {e}") from e


def extract_text_from_docx(file_path: Path):
    try:
        doc = DocxDocument(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        return [
            Document(
                page_content=full_text,
                metadata={
                    "source": file_path.name,
                    "page": 1
                }
            )
        ]
    except Exception as e:
        raise DocumentLoadError(f"DOCX extraction failed: {e}") from e


def load_document(file_path):
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    elif suffix == ".docx":
        return extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")